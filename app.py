from __future__ import annotations

import csv
import html
import io
import re
from pathlib import Path

import streamlit as st  # type: ignore[import-not-found]

from resume_analyzer import (
    TRAINED_MODEL_PATH,
    JobPosting,
    analyze_resume,
    format_report,
    rank_jobs_for_resume,
    trained_model_info,
)


DATASET_PATH = Path("resume_job_matching_dataset.csv")
DEFAULT_RESUME_FOLDER = Path("/Users/jigyasaverma/Desktop/Resume_screening/pythonProject/merged_resumes")
DESCRIPTION_COLUMNS = ("job_description", "job description", "description", "job desc", "jd", "details", "posting")
TITLE_COLUMNS = ("title", "job_title", "job title", "role", "position", "designation")
COMPANY_COLUMNS = ("company", "company_name", "company name", "organization", "employer")
URL_COLUMNS = ("url", "link", "job_url", "job url", "posting_url", "posting url")


def render_global_styles() -> None:
    st.markdown(
        """
        <style>
        .block-container { padding-top: 2rem; max-width: 1180px; }
        .hero {
            padding: 1.5rem 1.6rem;
            border: 1px solid #e8eaf0;
            border-radius: 18px;
            background: linear-gradient(135deg, #f8fbff 0%, #f7f4ff 100%);
            margin-bottom: 1.2rem;
        }
        .hero h1 { margin-bottom: 0.25rem; }
        .chip {
            display: inline-block;
            padding: 0.25rem 0.55rem;
            border-radius: 999px;
            margin: 0.16rem 0.2rem 0.16rem 0;
            font-size: 0.84rem;
            border: 1px solid transparent;
        }
        .chip-good { background: #e8f7ee; color: #116b37; border-color: #bfe8cd; }
        .chip-bad { background: #fff0f0; color: #9b1c1c; border-color: #ffd0d0; }
        .chip-warn { background: #fff7df; color: #8a5a00; border-color: #ffe3a1; }
        .muted-card {
            border: 1px solid #e8eaf0;
            border-radius: 14px;
            padding: 1rem;
            background: #ffffff;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_hero(title: str, subtitle: str) -> None:
    st.markdown(
        f"""
        <div class="hero">
            <h1>{html.escape(title)}</h1>
            <p>{html.escape(subtitle)}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_score_cards(result: dict[str, object]) -> None:
    bias = result["bias"] if isinstance(result.get("bias"), dict) else {}
    method = {
        "learning_to_rank": "Learning-to-rank",
        "trained_model": "Trained model",
    }.get(str(result.get("scoring_method")), "Formula fallback")
    col_final, col_semantic, col_ats, col_bias = st.columns(4)
    col_final.metric("Final Score", f"{result['final_score']}/100", method)
    col_semantic.metric("Semantic Match", f"{result['semantic_score']}%")
    col_ats.metric("ATS Score", f"{result['ats_score']}/100")
    col_bias.metric("Bias Risk", str(bias.get("risk_level", "Unknown")))


def render_skill_chips(title: str, skills: list[str], variant: str) -> None:
    if not skills:
        st.caption(f"{title}: none detected")
        return
    chips = " ".join(
        f'<span class="chip chip-{variant}">{html.escape(skill)}</span>'
        for skill in skills
    )
    st.markdown(f"**{html.escape(title)}**<br>{chips}", unsafe_allow_html=True)


def render_analysis_summary(result: dict[str, object]) -> None:
    skill_match = result["skill_match"]
    render_score_cards(result)
    st.divider()
    skill_cols = st.columns(3)
    with skill_cols[0]:
        render_skill_chips("Matched Skills", skill_match.matched, "good")
    with skill_cols[1]:
        render_skill_chips("Missing Skills", skill_match.missing, "bad")
    with skill_cols[2]:
        render_skill_chips("Partial Matches", skill_match.partial, "warn")

    st.markdown("### Explanation")
    st.write(result["overall_explanation"])

    with st.expander("Section-wise Resume Feedback", expanded=True):
        sections = result["sections"]
        for section_name in ["education", "experience", "projects", "skills"]:
            section = sections[section_name]
            st.markdown(f"**{section_name.title()}: {section['rating']}**")
            st.caption(section["feedback"])

    with st.expander("ATS and Improvement Suggestions"):
        st.markdown("**ATS suggestions**")
        for suggestion in result["ats_suggestions"]:
            st.write(f"- {suggestion}")
        st.markdown("**Resume improvements**")
        for suggestion in result["improvements"]:
            st.write(f"- {suggestion}")

    with st.expander("Full Structured Report"):
        st.markdown(format_report(result))


def extract_uploaded_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]

            reader = PdfReader(uploaded_file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            st.warning(f"Could not read PDF file {uploaded_file.name}: {exc}")
            return ""

    if name.endswith(".docx"):
        try:
            from docx import Document  # type: ignore[import-not-found]

            document = Document(uploaded_file)
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception as exc:
            st.warning(f"Could not read DOCX file {uploaded_file.name}: {exc}")
            return ""

    st.warning(f"Unsupported file type: {uploaded_file.name}")
    return ""


def read_resume_file(path: Path) -> str:
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if path.suffix.lower() == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    if path.suffix.lower() == ".docx":
        try:
            from docx import Document  # type: ignore[import-not-found]

            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return ""

    return ""


def iter_resume_files(folder_path: Path, max_files: int) -> list[Path]:
    files = []
    for path in folder_path.rglob("*"):
        if path.suffix.lower() in {".txt", ".pdf", ".docx"}:
            files.append(path)
        if len(files) >= max_files:
            break
    return files


def load_dataset_rows(limit: int = 25) -> list[dict[str, str]]:
    if not DATASET_PATH.exists():
        return []

    rows = []
    with DATASET_PATH.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for index, row in enumerate(reader):
            if index >= limit:
                break
            rows.append(row)
    return rows


def load_job_postings_from_dataset(limit: int) -> list[JobPosting]:
    rows = load_dataset_rows(limit)
    return [
        JobPosting(
            title=f"Dataset Job {index}",
            company="Dataset",
            description=row.get("job_description", ""),
        )
        for index, row in enumerate(rows, start=1)
        if row.get("job_description", "").strip()
    ]


def parse_jobs_csv(uploaded_file) -> tuple[list[JobPosting], str | None]:
    raw_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
    reader = csv.DictReader(io.StringIO(raw_text))
    if not reader.fieldnames:
        return [], "The CSV has no header row."

    normalized_fields = {_normalize_column_name(field): field for field in reader.fieldnames}
    description_column = _first_available_column(normalized_fields, DESCRIPTION_COLUMNS)
    if description_column is None:
        return [], "Could not find a job description column. Use one of: job_description, description, jd, details."

    title_column = _first_available_column(normalized_fields, TITLE_COLUMNS)
    company_column = _first_available_column(normalized_fields, COMPANY_COLUMNS)
    url_column = _first_available_column(normalized_fields, URL_COLUMNS)

    postings = []
    for index, row in enumerate(reader, start=1):
        description = row.get(description_column, "").strip()
        if not description:
            continue
        postings.append(
            JobPosting(
                description=description,
                title=row.get(title_column, "").strip() if title_column else f"Uploaded Job {index}",
                company=row.get(company_column, "").strip() if company_column else "",
                url=row.get(url_column, "").strip() if url_column else "",
            )
        )

    return postings, None


def _first_available_column(field_lookup: dict[str, str], candidates: tuple[str, ...]) -> str | None:
    for candidate in candidates:
        normalized_candidate = _normalize_column_name(candidate)
        if normalized_candidate in field_lookup:
            return field_lookup[normalized_candidate]
    return None


def _normalize_column_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.strip().lower()).strip("_")


def resolve_resume_text(uploaded_resume, pasted_text: str) -> tuple[str, str | None]:
    if uploaded_resume is not None:
        extracted = extract_uploaded_text(uploaded_resume)
        if extracted.strip():
            return extracted, f"Using text extracted from {uploaded_resume.name}."
    return pasted_text, None


def render_single_analyzer() -> None:
    st.subheader("Analyze Resume Against One Job")
    st.caption("Upload a resume, paste a job description, and get an explainable fit score.")
    dataset_rows = load_dataset_rows()
    default_resume = ""
    default_job = ""

    with st.expander("Use a sample from the dataset"):
        sample_label = "Use custom text"
        if dataset_rows:
            sample_options = ["Use custom text"] + [
                f"Sample {index + 1}: score {row.get('match_score', 'N/A')}"
                for index, row in enumerate(dataset_rows)
            ]
            sample_label = st.selectbox("Dataset sample", sample_options)
        if dataset_rows and sample_label != "Use custom text":
            sample_index = int(sample_label.split(":")[0].replace("Sample", "").strip()) - 1
            default_resume = dataset_rows[sample_index].get("resume", "")
            default_job = dataset_rows[sample_index].get("job_description", "")

    upload_col, job_col = st.columns([1, 1])
    with upload_col:
        st.markdown("#### 1. Resume")
        uploaded_resume = st.file_uploader("Upload resume", type=["txt", "pdf", "docx"])
        resume_text = st.text_area("Or paste resume text", value=default_resume, height=210)
    with job_col:
        st.markdown("#### 2. Job Description")
        job_text = st.text_area("Paste job description", value=default_job, height=280)

    if uploaded_resume is not None:
        resume_text, notice = resolve_resume_text(uploaded_resume, resume_text)
        if notice:
            st.info(notice)

    if resume_text.strip():
        with st.expander("Preview extracted resume text"):
            st.write(resume_text[:2500] + ("..." if len(resume_text) > 2500 else ""))

    if st.button("Analyze Match", type="primary", use_container_width=True):
        if not resume_text.strip() or not job_text.strip():
            st.error("Please provide both resume text and job description text.")
            return

        with st.spinner("Analyzing semantic match, skills, ATS, and bias..."):
            result = analyze_resume(resume_text, job_text)
        render_analysis_summary(result)


def render_job_recommendations() -> None:
    st.subheader("Find Best Matching Jobs")
    st.write(
        "Upload one resume, then rank many job descriptions from the included dataset or your own CSV export."
    )

    uploaded_resume = st.file_uploader(
        "Upload resume for job recommendations (.txt, .pdf, .docx)",
        type=["txt", "pdf", "docx"],
        key="recommend_resume",
    )
    pasted_resume = st.text_area("Or Paste Resume Text", height=240, key="recommend_resume_text")
    resume_text, notice = resolve_resume_text(uploaded_resume, pasted_resume)
    if notice:
        st.info(notice)

    source = st.radio(
        "Job description source",
        ["Use included dataset", "Upload jobs CSV"],
        horizontal=True,
    )
    max_jobs = st.number_input("Maximum jobs to compare", min_value=1, max_value=200, value=25, key="recommend_max")
    top_n = st.number_input("Show top matches", min_value=1, max_value=50, value=10, key="recommend_top")

    postings: list[JobPosting] = []
    if source == "Use included dataset":
        postings = load_job_postings_from_dataset(int(max_jobs))
        st.caption("Using `resume_job_matching_dataset.csv` as the job source.")
    else:
        uploaded_jobs = st.file_uploader(
            "Upload jobs CSV",
            type=["csv"],
            key="jobs_csv",
            help="Expected description column: job_description, description, jd, details, or posting.",
        )
        if uploaded_jobs is not None:
            postings, error = parse_jobs_csv(uploaded_jobs)
            if error:
                st.error(error)
                postings = []
            else:
                postings = postings[: int(max_jobs)]
                st.success(f"Loaded {len(postings)} job descriptions from CSV.")

    if st.button("Rank Matching Jobs", type="primary"):
        if not resume_text.strip():
            st.error("Please upload or paste a resume first.")
            return
        if not postings:
            st.error("No job descriptions are available to rank.")
            return

        with st.spinner("Ranking jobs and generating explainable match reports..."):
            ranked = rank_jobs_for_resume(resume_text, postings, limit=int(top_n))

        if not ranked:
            st.warning("No usable job descriptions were found.")
            return

        table_rows = [
            {
                "rank": item["rank"],
                "title": item["title"],
                "company": item["company"],
                "final_score": item["final_score"],
                "scoring_method": item["analysis"]["scoring_method"],
                "semantic_score": item["semantic_score"],
                "skill_score": item["skill_score"],
                "ats_score": item["ats_score"],
                "bias_risk": item["bias_risk"],
                "matched_skills": ", ".join(item["matched_skills"]),
                "missing_skills": ", ".join(item["missing_skills"]),
                "url": item["url"],
            }
            for item in ranked
        ]
        st.dataframe(table_rows, use_container_width=True, hide_index=True)

        st.markdown("### Explainable Match Details")
        for item in ranked:
            label = f"#{item['rank']} {item['title']} - {item['final_score']}/100"
            with st.expander(label):
                if item["company"]:
                    st.write(f"Company: {item['company']}")
                if item["url"]:
                    st.write(f"URL: {item['url']}")
                st.write(item["summary"])
                st.markdown(format_report(item["analysis"]))


def render_folder_analyzer() -> None:
    st.subheader("Analyze Resume Folder")
    folder_path = Path(st.text_input("Resume folder path", value=str(DEFAULT_RESUME_FOLDER)))
    job_text = st.text_area("Job Description Text for Folder Analysis", height=220, key="folder_job")
    max_files = st.number_input("Maximum files to analyze", min_value=1, max_value=200, value=25)

    if st.button("Analyze Folder"):
        if not folder_path.exists():
            st.error("Folder path does not exist.")
            return
        if not job_text.strip():
            st.error("Please provide a job description.")
            return

        files = iter_resume_files(folder_path, int(max_files))

        if not files:
            st.warning("No .txt, .pdf, or .docx resumes found in the selected folder.")
            return

        rows = []
        progress = st.progress(0)
        for index, path in enumerate(files, start=1):
            resume_text = read_resume_file(path)
            if resume_text.strip():
                result = analyze_resume(resume_text, job_text)
                rows.append(
                    {
                        "file": path.name,
                        "final_score": result["final_score"],
                        "semantic_score": result["semantic_score"],
                        "ats_score": result["ats_score"],
                        "bias_risk": result["bias"]["risk_level"],
                    }
                )
            progress.progress(index / len(files))

        if rows:
            rows.sort(key=lambda item: item["final_score"], reverse=True)
            st.dataframe(rows, use_container_width=True)
        else:
            st.warning("Files were found, but text could not be extracted from them.")


def render_model_training() -> None:
    st.subheader("Model Training")
    st.write("Train a supervised model from `resume_job_matching_dataset.csv` so final scores learn from your dataset.")

    info = trained_model_info()
    if info["available"]:
        st.success("Trained model found. New analyses will use it automatically.")
        st.write(f"Model path: `{info['path']}`")
        st.write(f"Rows used for training: `{info.get('trained_rows')}`")
        st.write(f"Target scale: `{info.get('target_scale')}`")
        st.write(f"Model type: `{info.get('model_type')}`")
        metrics = info.get("metrics", {})
        if isinstance(metrics, dict) and metrics:
            st.write(
                {
                    "MAE": metrics.get("mae"),
                    "RMSE": metrics.get("rmse"),
                    "R2": metrics.get("r2"),
                    "Test rows": metrics.get("test_rows"),
                }
            )
        rank_metrics = info.get("rank_metrics", {})
        if isinstance(rank_metrics, dict) and rank_metrics.get("available"):
            st.write(
                {
                    "Ranker pairs": rank_metrics.get("pair_count"),
                    "Ranker accuracy": rank_metrics.get("accuracy"),
                    "Ranker test pairs": rank_metrics.get("test_pairs"),
                }
            )
    else:
        st.warning("No trained model found yet. The analyzer is currently using the formula fallback.")
        st.write(f"Expected model path: `{TRAINED_MODEL_PATH}`")

    st.markdown("Run this command in the project folder to train or retrain the model:")
    st.code("python3 train_match_model.py --dataset resume_job_matching_dataset.csv", language="bash")
    st.caption("After training, restart Streamlit so the app loads the saved model artifact.")

def main() -> None:
    st.set_page_config(page_title="AI Resume Analyzer", layout="wide")
    render_global_styles()
    render_hero(
        "AI-Powered Resume Analyzer",
        "Upload resumes, score job fit, rank opportunities, and get explainable improvement guidance.",
    )

    tab_single, tab_recommend, tab_folder, tab_training = st.tabs(
        ["Single Resume", "Job Recommendations", "Resume Folder", "Model Training"]
    )
    with tab_single:
        render_single_analyzer()
    with tab_recommend:
        render_job_recommendations()
    with tab_folder:
        render_folder_analyzer()
    with tab_training:
        render_model_training()


if __name__ == "__main__":
    main()
